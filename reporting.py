import io

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from config import INTERNAL_TEMP_KEYS, INTERNAL_HUMIDITY_KEYS, EXTERNAL_SURFACE_KEYS, INTERNAL_SURFACE_KEYS
from plotting import build_plot_figure


def fig_to_png_bytes(fig):
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=200, bbox_inches="tight")
    buffer.seek(0)
    plt.close(fig)
    return buffer


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    return p


def add_table_from_df(doc, df):
    display_df = df.copy()
    for col in display_df.columns:
        if pd.api.types.is_numeric_dtype(display_df[col]):
            display_df[col] = display_df[col].map(lambda x: "" if pd.isna(x) else f"{x:.2f}" if isinstance(x, float) else str(x))
        else:
            display_df[col] = display_df[col].fillna("").astype(str)

    table = doc.add_table(rows=1, cols=len(display_df.columns))
    table.style = "Table Grid"
    for i, col_name in enumerate(display_df.columns):
        table.rows[0].cells[i].text = str(col_name)

    for _, row in display_df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def add_plot_group(doc, title, keys, data_map, plot_prefix, all_stats, day_stats):
    add_heading(doc, title, level=1)
    for key in keys:
        if key not in data_map:
            continue
        add_heading(doc, key, level=3)
        fig = build_plot_figure(data_map[key], f"{plot_prefix} - Full Day - {key}")
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.5))

    add_heading(doc, "full day", level=2)
    add_table_from_df(doc, all_stats[all_stats["Position"].isin(keys)])

    add_heading(doc, "day time", level=2)
    add_table_from_df(doc, day_stats[day_stats["Position"].isin(keys)])


def generate_detailed_report(
    start_date,
    end_date,
    internal_allday_map,
    internal_daytime_map,
    internal_allday_stats,
    internal_daytime_stats,
    sheds_allday_map,
    sheds_daytime_map,
    sheds_allday_stats,
    sheds_daytime_stats,
):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Detailed report")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Reporting window: {start_date} to {end_date}").italic = True

    add_plot_group(doc, "Internal Temperature of the sheds", INTERNAL_TEMP_KEYS, internal_allday_map, "Internal Temperature", internal_allday_stats, internal_daytime_stats)
    add_plot_group(doc, "Internal Humidity of the sheds", INTERNAL_HUMIDITY_KEYS, internal_allday_map, "Internal Humidity", internal_allday_stats, internal_daytime_stats)
    add_plot_group(doc, "External Surface Temperature of the sheds", EXTERNAL_SURFACE_KEYS, sheds_allday_map, "External Surface Temperature", sheds_allday_stats, sheds_daytime_stats)
    add_plot_group(doc, "Internal Surface Temperature of the sheds", INTERNAL_SURFACE_KEYS, sheds_allday_map, "Internal Surface Temperature", sheds_allday_stats, sheds_daytime_stats)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
