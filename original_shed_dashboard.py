
import io
from datetime import time
from pydoc import doc

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

st.set_page_config(page_title="Shed Data Dashboard", layout="wide")
st.title("Shed Data Dashboard")
st.caption("Unified web app for Internal Reading, Sheds Reading, and Weather Station analysis")


# =============================
# Shared helpers
# =============================
def read_excel_with_time_index(uploaded_file):
    if uploaded_file is None:
        return None
    df = pd.read_excel(uploaded_file, index_col="Time")
    df.index = pd.to_datetime(df.index)
    return df


def filter_by_daterange(df, start_date, end_date):
    start_dt = pd.Timestamp(start_date)
    end_dt = pd.Timestamp(end_date) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)
    return df[(df.index >= start_dt) & (df.index <= end_dt)]


def filter_by_timewindow(df, start_date, end_date, start_time, end_time):
    filtered = filter_by_daterange(df, start_date, end_date)
    return filtered.between_time(start_time.strftime("%H:%M"), end_time.strftime("%H:%M"))


def build_group_map(df, group_map):
    return {key: df[columns].copy() for key, columns in group_map.items()}


def sync_columns(cols_df):
    return cols_df.dropna(how="any")


def clean_temperature_data(cols_df, t_min, t_max, t_change):
    cleaned = cols_df.copy()

    rows_to_remove = pd.Series(False, index=cleaned.index)

    out_of_range = (cleaned < t_min) | (cleaned > t_max)
    rows_to_remove |= out_of_range.any(axis=1)

    diffs = cleaned.diff().abs()
    spike_mask = diffs > t_change
    rows_to_remove |= spike_mask.any(axis=1)

    cleaned.loc[rows_to_remove, :] = np.nan
    return sync_columns(cleaned)


def process_map(column_map, start_date, end_date, start_time, end_time, clean_fn=None):
    processed_map = {k: clean_fn(v) for k, v in column_map.items()} if clean_fn else column_map

    allday_map = {
        k: filter_by_daterange(v, start_date, end_date)
        for k, v in processed_map.items()
    }

    daytime_map = {
        k: filter_by_timewindow(v, start_date, end_date, start_time, end_time)
        for k, v in processed_map.items()
    }

    allday_stats = stats_from_map(allday_map)
    daytime_stats = stats_from_map(daytime_map)

    return processed_map, allday_map, daytime_map, allday_stats, daytime_stats


def stats_from_map(data_map):
    rows = []
    for key, cols in data_map.items():
        means = cols.mean()
        stds = cols.std()
        for col_name in cols.columns:
            rows.append(
                {
                    "Position": key,
                    "Series": col_name,
                    "Mean": round(float(means[col_name]), 2) if pd.notna(means[col_name]) else np.nan,
                    "Std Dev": round(float(stds[col_name]), 2) if pd.notna(stds[col_name]) else np.nan,
                    "Count": int(cols[col_name].count()),
                }
            )
    return pd.DataFrame(rows)


def summarize_selected_series(data_map, selections, prefix):
    summary = {}

    for shed, column_names in selections.items():
        values = []
        for group_key, col_name in column_names:
            if group_key in data_map and col_name in data_map[group_key].columns:
                values.extend(data_map[group_key][col_name].dropna().tolist())

        if len(values) > 0:
            summary[shed] = {
                "Mean": round(np.mean(values), 2),
                "Std Dev": round(np.std(values, ddof=1), 2) if len(values) > 1 else 0.0,
            }
        else:
            summary[shed] = {
                "Mean": np.nan,
                "Std Dev": np.nan,
            }

    summary_df = pd.DataFrame(summary).T.reset_index()
    summary_df.columns = ["Shed", f"{prefix} Mean", f"{prefix} Std Dev"]
    return summary_df


def make_internal_summary(map_data, sensor_prefix, prefix_label):
    selections = {
        shed: [(f"{sensor_prefix}{i}", f"{shed}-{sensor_prefix}{i}") for i in range(1, 5)]
        for shed in ["R", "C", "TII"]
    }
    return summarize_selected_series(map_data, selections, prefix_label)


def make_external_summary(map_data, positions, prefix_label):
    selections = {
        shed: [
            (pos, f"{shed}-TCO{pos.replace('O', '')}")
            for pos in positions
        ]
        for shed in ["R", "C", "TII"]
    }
    return summarize_selected_series(map_data, selections, prefix_label)


def make_internal_surface_summary(map_data, positions, prefix_label):
    selections = {
        shed: [
            (pos, f"{shed}-TCI{pos.replace('I', '')}")
            for pos in positions
        ]
        for shed in ["R", "C", "TII"]
    }
    return summarize_selected_series(map_data, selections, prefix_label)


def plot_map(data_map, title_prefix):
    for key, cols in data_map.items():
        fig = build_plot_figure(cols, f"{title_prefix} - {key}")
        st.pyplot(fig)


def build_plot_figure(cols, title):
    fig, ax = plt.subplots(figsize=(9, 4))
    for col in cols.columns:
        ax.scatter(cols.index, cols[col], s=6, label=col)
    ax.set_title(title, pad=18)
    ax.tick_params(axis="x", rotation=30)
    legend_columns = min(len(cols.columns), 3)
    ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.5, -0.25),
        ncol=legend_columns,
        frameon=False,
    )
    fig.tight_layout()
    return fig


def display_stats_pair(all_df, day_df):
    c1, c2 = st.columns(2)
    with c1:
        st.write("All day statistics")
        st.dataframe(all_df, use_container_width=True)
    with c2:
        st.write("Day window statistics")
        st.dataframe(day_df, use_container_width=True)


def export_excel(sheet_dict):
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for sheet_name, df in sheet_dict.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    buffer.seek(0)
    return buffer


def download_excel_button(label, file_name, sheet_dict):
    excel_data = export_excel(sheet_dict)
    st.download_button(
        label,
        data=excel_data,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def fig_to_png_bytes(fig):
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=200, bbox_inches="tight")
    buffer.seek(0)
    plt.close(fig)
    return buffer


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    return p


def add_table_from_df(doc, df):
    display_df = df.copy()
    for col in display_df.columns:
        if pd.api.types.is_numeric_dtype(display_df[col]):
            display_df[col] = display_df[col].map(
                lambda x: "" if pd.isna(x) else f"{x:.2f}" if isinstance(x, float) else str(x)
            )
        else:
            display_df[col] = display_df[col].fillna("").astype(str)

    table = doc.add_table(rows=1, cols=len(display_df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(display_df.columns):
        hdr_cells[i].text = str(col_name)

    for _, row in display_df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def add_plot_section(doc, title, keys, data_map, plot_title_prefix, table_df, heading_level=1, key_heading_level=2):
    add_heading(doc, title, level=heading_level)
    for key in keys:
        add_heading(doc, key, level=key_heading_level)
        fig = build_plot_figure(data_map[key], f"{plot_title_prefix} - {key}")
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.5))
    add_table_from_df(doc, table_df)

def clean_solar_data(df, col_name, solar_min=-50, solar_max=1400, max_step=300):
    cleaned = df.copy()

    # remove unrealistically low values
    cleaned.loc[cleaned[col_name] < solar_min, col_name] = np.nan

    # remove unrealistically high values
    cleaned.loc[cleaned[col_name] > solar_max, col_name] = np.nan

    # remove sudden spikes
    diffs = cleaned[col_name].diff().abs()
    cleaned.loc[diffs > max_step, col_name] = np.nan

    # drop rows where solar value became NaN
    cleaned = cleaned.dropna(subset=[col_name])

    return cleaned

def calculate_delta_summary(summary_df, mean_col, selected_deltas):
    means = summary_df.set_index("Shed")[mean_col]
    results = []

    if "ΔT1" in selected_deltas and "TII" in means.index and "C" in means.index:
        results.append({"Delta": "ΔT1", "Formula": "TII - C", "Value": means["TII"] - means["C"]})

    if "ΔT2" in selected_deltas and "TII" in means.index and "R" in means.index:
        results.append({"Delta": "ΔT2", "Formula": "TII - R", "Value": means["TII"] - means["R"]})

    if "ΔT3" in selected_deltas and "C" in means.index and "R" in means.index:
        results.append({"Delta": "ΔT3", "Formula": "C - R", "Value": means["C"] - means["R"]})

    return pd.DataFrame(results)

def export_excel_grouped(grouped_dict):
    buffer = io.BytesIO()

    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        for sheet_name, sections in grouped_dict.items():
            startrow = 0

            for title, df in sections:
                # write section title
                pd.DataFrame([[title]]).to_excel(
                    writer,
                    sheet_name=sheet_name,
                    startrow=startrow,
                    index=False,
                    header=False
                )

                startrow += 1

                # write dataframe
                df.to_excel(
                    writer,
                    sheet_name=sheet_name,
                    startrow=startrow,
                    index=False
                )

                startrow += len(df) + 3  # spacing between tables

    buffer.seek(0)
    return buffer

def generate_detailed_report(
    start_date,
    end_date,
    internal_allday_map,
    internal_daytime_map,
    internal_allday_stats,
    internal_daytime_stats,
    sheds_allday_map,
    sheds_daytime_map,
    sheds_allday_stats,
    sheds_daytime_stats,
):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Detailed report")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Reporting window: {start_date} to {end_date}").italic = True

    # Internal Temperature
    add_heading(doc, "Internal Temperature of the sheds", level=1)
    for key in ["T1", "T2", "T3", "T4"]:
        add_heading(doc, key, level=3)
        fig = build_plot_figure(internal_allday_map[key], f"Internal Temperature - Full Day - {key}")
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.5))
    add_heading(doc, "full day", level=2)
    add_table_from_df(
        doc,
        internal_allday_stats[internal_allday_stats["Position"].isin(["T1", "T2", "T3", "T4"])]
    )
    add_heading(doc, "day time", level=2)
    add_table_from_df(
        doc,
        internal_daytime_stats[internal_daytime_stats["Position"].isin(["T1", "T2", "T3", "T4"])]
    )

    # Internal Humidity
    add_heading(doc, "Internal Humidity of the sheds", level=1)
    for key in ["H1", "H2", "H3", "H4"]:
        add_heading(doc, key, level=3)
        fig = build_plot_figure(internal_allday_map[key], f"Internal Humidity - Full Day - {key}")
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.5))
    add_heading(doc, "full day", level=2)
    add_table_from_df(
    doc,
    internal_allday_stats[internal_allday_stats["Position"].isin(["H1", "H2", "H3", "H4"])]
)
    add_heading(doc, "day time", level=2)
    add_table_from_df(
        doc,
        internal_daytime_stats[internal_daytime_stats["Position"].isin(["H1", "H2", "H3", "H4"])]
    )

    # External Surface
    add_heading(doc, "External Surface Temperature of the sheds", level=1)
    for key in ["O1", "O2", "O3", "O4", "O5"]:
        add_heading(doc, key, level=3)
        fig = build_plot_figure(sheds_allday_map[key], f"External Surface Temperature - Full Day - {key}")
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.5))
    add_heading(doc, "full day", level=2)
    add_table_from_df(doc, sheds_allday_stats[sheds_allday_stats["Position"].isin(["O1", "O2", "O3", "O4", "O5"])])
    add_heading(doc, "day time", level=2)
    add_table_from_df(doc, sheds_daytime_stats[sheds_daytime_stats["Position"].isin(["O1", "O2", "O3", "O4", "O5"])])

    # Internal Surface
    add_heading(doc, "Internal Surface Temperature of the sheds", level=1)
    for key in ["I1", "I2", "I3", "I4", "I5", "I6"]:
        add_heading(doc, key, level=3)
        fig = build_plot_figure(sheds_allday_map[key], f"Internal Surface Temperature - Full Day - {key}")
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6.5))
    add_heading(doc, "full day", level=2)
    add_table_from_df(doc, sheds_allday_stats[sheds_allday_stats["Position"].isin(["I1", "I2", "I3", "I4", "I5", "I6"])])
    add_heading(doc, "day time", level=2)
    add_table_from_df(doc, sheds_daytime_stats[sheds_daytime_stats["Position"].isin(["I1", "I2", "I3", "I4", "I5", "I6"])])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================
# Sidebar controls
# =============================
st.sidebar.header("Controls")
start_date = st.sidebar.date_input("Start date")
end_date = st.sidebar.date_input("End date")
start_time = st.sidebar.time_input("Day window start", value=time(9, 30))
end_time = st.sidebar.time_input("Day window end", value=time(15, 30))
show_raw = st.sidebar.checkbox("Show raw plots", value=False)

if start_date > end_date:
    st.error("Start date must be on or before end date.")
    st.stop()


# =============================
# Uploads and common inputs
# =============================
tab1, tab2, tab3 = st.tabs(["Detailed Data", "Weather Station", "Processed Data"])

with tab1:
    st.subheader("Data Upload")
    internal_file = st.file_uploader("Upload internal sheds file", type=["xlsx", "xls"], key="internal")
    sheds_file = st.file_uploader("Upload sheds probe file", type=["xlsx", "xls"], key="sheds")

    t_min = st.number_input("Minimum valid temperature", value=0.0, key="tmin")
    t_max = st.number_input("Maximum valid temperature", value=75.0, key="tmax")
    t_change = st.number_input("Spike threshold", value=5.0, key="tchange")


# =============================
# Shared data preparation
# =============================
INTERNAL_GROUPS = {
    "T1": ["R-T1", "C-T1", "TII-T1"],
    "T2": ["R-T2", "C-T2", "TII-T2"],
    "T3": ["R-T3", "C-T3", "TII-T3"],
    "T4": ["R-T4", "C-T4", "TII-T4"],
    "H1": ["R-H1", "C-H1", "TII-H1"],
    "H2": ["R-H2", "C-H2", "TII-H2"],
    "H3": ["R-H3", "C-H3", "TII-H3"],
    "H4": ["R-H4", "C-H4", "TII-H4"],
}

SHEDS_GROUPS = {
    "O1": ["R-TCO1", "C-TCO1", "TII-TCO1"],
    "O2": ["R-TCO2", "C-TCO2", "TII-TCO2"],
    "O3": ["R-TCO3", "C-TCO3", "TII-TCO3"],
    "O4": ["R-TCO4", "C-TCO4", "TII-TCO4"],
    "O5": ["R-TCO5", "C-TCO5", "TII-TCO5"],
    "I1": ["R-TCI1", "C-TCI1", "TII-TCI1"],
    "I2": ["R-TCI2", "C-TCI2", "TII-TCI2"],
    "I3": ["R-TCI3", "C-TCI3", "TII-TCI3"],
    "I4": ["R-TCI4", "C-TCI4", "TII-TCI4"],
    "I5": ["R-TCI5", "C-TCI5", "TII-TCI5"],
    "I6": ["R-TCI6", "C-TCI6", "TII-TCI6"],
}

df_internal = read_excel_with_time_index(internal_file) if "internal_file" in locals() else None
df_sheds = read_excel_with_time_index(sheds_file) if "sheds_file" in locals() else None

internal_column_map = build_group_map(df_internal, INTERNAL_GROUPS) if df_internal is not None else None
sheds_column_map = build_group_map(df_sheds, SHEDS_GROUPS) if df_sheds is not None else None

internal_processed_map = internal_allday_map = internal_daytime_map = None
internal_allday_stats = internal_daytime_stats = None

sheds_cleaned_map = sheds_allday_map = sheds_daytime_map = None
sheds_allday_stats = sheds_daytime_stats = None

if internal_column_map is not None:
    (
        internal_processed_map,
        internal_allday_map,
        internal_daytime_map,
        internal_allday_stats,
        internal_daytime_stats,
    ) = process_map(
        internal_column_map,
        start_date,
        end_date,
        start_time,
        end_time,
    )

if sheds_column_map is not None:
    (
        sheds_cleaned_map,
        sheds_allday_map,
        sheds_daytime_map,
        sheds_allday_stats,
        sheds_daytime_stats,
    ) = process_map(
        sheds_column_map,
        start_date,
        end_date,
        start_time,
        end_time,
        clean_fn=lambda df: clean_temperature_data(df, t_min, t_max, t_change),
    )


# =============================
# Tab 1: Detailed Data
# =============================
with tab1:
    if internal_column_map is not None:
        st.markdown("### Internal Reading")
        if show_raw:
            plot_map(internal_column_map, "Raw")
        plot_map(internal_allday_map, "Filtered")
        display_stats_pair(internal_allday_stats, internal_daytime_stats)

        excel_data1 = export_excel_grouped({
            "Internal": [
                ("All Day", internal_allday_stats),
                ("Day Window", internal_daytime_stats),
            ],
        })

        st.download_button(
            "Download internal Excel report",
            data=excel_data1,
            file_name=f"internal_stats_{start_date}_to_{end_date}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    if sheds_column_map is not None:
        st.markdown("### Sheds Reading")
        if show_raw:
            plot_map(sheds_column_map, "Raw")
        plot_map(sheds_allday_map, "Filtered")
        display_stats_pair(sheds_allday_stats, sheds_daytime_stats)

        excel_data2 = export_excel_grouped({
            "Sheds": [
                ("All Day", sheds_allday_stats),
                ("Day Window", sheds_daytime_stats),
            ],
        })

        st.download_button(
            "Download sheds surface Excel report",
            data=excel_data2,
            file_name=f"sheds_stats_{start_date}_to_{end_date}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    if internal_column_map is not None and sheds_column_map is not None:
        st.markdown("### Detailed Word Report")
        report_buffer = generate_detailed_report(
            start_date,
            end_date,
            internal_allday_map,
            internal_daytime_map,
            internal_allday_stats,
            internal_daytime_stats,
            sheds_allday_map,
            sheds_daytime_map,
            sheds_allday_stats,
            sheds_daytime_stats,
        )
        st.download_button(
            "Download detailed Word report",
            data=report_buffer,
            file_name=f"detailed_report_{start_date}_to_{end_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    elif internal_column_map is not None or sheds_column_map is not None:
        st.info("Upload both the internal sheds file and the sheds probe file to enable the detailed Word report button.")


# =============================
# Tab 2: Weather Station
# =============================
with tab2:
    st.subheader("Data Upload")
    weather_file = st.file_uploader("Upload weather station file", type=["xlsx", "xls"], key="weather")
    solar_file = st.file_uploader("Upload solar irradiation file", type=["xlsx", "xls"], key="solar")

    # optional controls for solar cleaning
    solar_min = st.number_input("Minimum valid solar irradiance", value=-50.0, key="solar_min")
    solar_max = st.number_input("Maximum valid solar irradiance", value=1400.0, key="solar_max")
    solar_change = st.number_input("Solar spike threshold", value=300.0, key="solar_change")

    if weather_file is not None and solar_file is not None:
        df_weather = pd.read_excel(weather_file, index_col="Time", parse_dates=True)
        df_solar = pd.read_excel(solar_file, index_col="Time", parse_dates=True)

        df_weather.index = pd.to_datetime(df_weather.index)
        df_solar.index = pd.to_datetime(df_solar.index)
        df_solar = df_solar.apply(pd.to_numeric, errors="coerce")

        # -----------------------------
        # WEATHER: temperature + humidity
        # -----------------------------
        weather_cols = df_weather[["Outdoor Temperature(ºC)", "Outdoor Humidity(%RH)"]].copy()

        weather_allday = filter_by_daterange(weather_cols, start_date, end_date)
        weather_daytime = filter_by_timewindow(weather_cols, start_date, end_date, start_time, end_time)

        def summarize_weather(df):
            rows = []
            for col in ["Outdoor Temperature(ºC)", "Outdoor Humidity(%RH)"]:
                series = df[col].dropna()
                rows.append({
                    "Parameter": col,
                    "Mean": round(series.mean(), 2) if not series.empty else np.nan,
                    "Std Dev": round(series.std(), 2) if len(series) > 1 else np.nan,
                    "Count": int(series.count())
                })
            return pd.DataFrame(rows)

        weather_allday_stats = summarize_weather(weather_allday)
        weather_daytime_stats = summarize_weather(weather_daytime)

        st.subheader("Outdoor Temperature and Humidity")
        c1, c2 = st.columns(2)
        with c1:
            st.write("All day statistics")
            st.dataframe(weather_allday_stats, use_container_width=True)
        with c2:
            st.write("Day window statistics")
            st.dataframe(weather_daytime_stats, use_container_width=True)

        fig_weather = build_plot_figure(weather_allday, "Outdoor Temperature and Humidity")
        st.pyplot(fig_weather)

        # -----------------------------
        # SOLAR: irradiance only
        # -----------------------------
        solar_col = "Smart sensor 1 Average"   # or "Smart sensor 1 Min."
        solar_df = df_solar[[solar_col]].copy()

        # filter by date first
        solar_allday = filter_by_daterange(solar_df, start_date, end_date)
        solar_daytime = filter_by_timewindow(solar_df, start_date, end_date, start_time, end_time)

        # show raw plot before cleaning if checkbox is selected
        if show_raw:
            st.write("Raw Solar Irradiance vs Date")
            fig_solar_raw = build_plot_figure(solar_allday, "Raw Solar Irradiance")
            st.pyplot(fig_solar_raw)

        # clean solar values
        solar_allday = clean_solar_data(
            solar_allday,
            solar_col,
            solar_min=solar_min,
            solar_max=solar_max,
            max_step=solar_change,
        )

        solar_daytime = clean_solar_data(
            solar_daytime,
            solar_col,
            solar_min=solar_min,
            solar_max=solar_max,
            max_step=solar_change,
        )

        st.subheader("Solar Irradiance")
        fig_solar = build_plot_figure(solar_allday, "Filtered Solar Irradiance")
        st.pyplot(fig_solar)

        # maximum irradiance from cleaned data
        solar_series = solar_allday[solar_col].dropna()
        if not solar_series.empty:
            max_irradiance = solar_series.max()
            max_time = solar_series.idxmax()

            st.success(
                f"Maximum irradiance in the selected range is {max_irradiance:.2f} "
                f"and it occurs on {max_time.strftime('%Y-%m-%d')} at {max_time.strftime('%H:%M:%S')}."
            )
        else:
            st.warning("No solar irradiance data available in the selected date range after filtering.")

        # -----------------------------
        # Excel export
        # -----------------------------
        excel_data3 = export_excel_grouped({
            "Weather": [
                ("All Day", weather_allday_stats),
                ("Day Window", weather_daytime_stats),
            ],
            "Solar": [
                ("All Day Filtered", solar_allday.reset_index()),
            ],
        })

        st.download_button(
            "Download weather station Excel report",
            data=excel_data3,
            file_name=f"weather_station_stats_{start_date}_to_{end_date}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

# =============================
# Tab 3: Processed Data
# =============================
with tab3:
    st.subheader("Important Information")
    st.info("O1 - out door, O2 - out right, O3 - out back, O4 - out left, O5 - out roof")
    st.info("I1 - in door, I2 - in right, I3 - in back, I4 - in left, I5 - in roof, I6 - in floor")

    if internal_allday_map is not None:
        st.subheader("Internal Reading")

        temp_all_summary = make_internal_summary(internal_allday_map, "T", "Temperature")
        temp_day_summary = make_internal_summary(internal_daytime_map, "T", "Temperature")
        hum_all_summary = make_internal_summary(internal_allday_map, "H", "Humidity")
        hum_day_summary = make_internal_summary(internal_daytime_map, "H", "Humidity")

        st.write("Average Internal Temperature (T1–T4)")
        display_stats_pair(temp_all_summary, temp_day_summary)

        st.write("Average Internal Humidity (H1–H4)")
        display_stats_pair(hum_all_summary, hum_day_summary)

        excel_data4 = export_excel_grouped({
            "Temperature": [
                ("All Day", temp_all_summary),
                ("Day Window", temp_day_summary),
            ],
            "Humidity": [
                ("All Day", hum_all_summary),
                ("Day Window", hum_day_summary),
            ],
        })

        st.download_button(
            "Download processed internal Excel report",
            data=excel_data4,
            file_name=f"processed_internal_stats_{start_date}_to_{end_date}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    if sheds_allday_map is not None:
        st.subheader("External Surface Reading")
        shaded_positions = ["O2", "O3"]
        nonshaded_positions = ["O4", "O5"]

        external_shaded_all_summary = make_external_summary(
            sheds_allday_map, shaded_positions, "Ext Surface T"
        )
        external_shaded_day_summary = make_external_summary(
            sheds_daytime_map, shaded_positions, "Ext Surface T"
        )

        external_nonshaded_all_summary = make_external_summary(
            sheds_allday_map, nonshaded_positions, "Ext Surface T"
        )
        external_nonshaded_day_summary = make_external_summary(
            sheds_daytime_map, nonshaded_positions, "Ext Surface T"
        )

        st.write("Average Shaded External Surface Temperature (O2–O3)")
        display_stats_pair(external_shaded_all_summary, external_shaded_day_summary)

        st.write("Average Non-Shaded External Surface Temperature (O4–O5)")
        display_stats_pair(external_nonshaded_all_summary, external_nonshaded_day_summary)

        st.subheader("Internal Surface Reading")

        internal_shaded_positions = ["I2", "I3"]
        internal_nonshaded_positions = ["I4", "I5"]
        internal_floor = ["I6"]

        internal_shaded_surface_all_summary = make_internal_surface_summary(
            sheds_allday_map, internal_shaded_positions, "Int Surface T"
        )
        internal_shaded_day_summary = make_internal_surface_summary(
            sheds_daytime_map, internal_shaded_positions, "Int Surface T"
        )

        internal_nonshaded_surface_all_summary = make_internal_surface_summary(
            sheds_allday_map, internal_nonshaded_positions, "Int Surface T"
        )
        internal_nonshaded_day_summary = make_internal_surface_summary(
            sheds_daytime_map, internal_nonshaded_positions, "Int Surface T"
        )

        internal_floor_all_summary = make_internal_surface_summary(
            sheds_allday_map, internal_floor, "Int Floor T"
        )
        internal_floor_day_summary = make_internal_surface_summary(
            sheds_daytime_map, internal_floor, "Int Floor T"
        )

        st.write("Average Shaded Internal Surface Temperature (I2–I3)")
        display_stats_pair(internal_shaded_surface_all_summary, internal_shaded_day_summary)

        st.write("Average Non-Shaded Internal Surface Temperature (I4–I5)")
        display_stats_pair(internal_nonshaded_surface_all_summary, internal_nonshaded_day_summary)

        st.write("Average Internal Floor Temperature (I6)")
        display_stats_pair(internal_floor_all_summary, internal_floor_day_summary)

        excel_data5 = export_excel_grouped({
            "External Shaded": [
                ("All Day", external_shaded_all_summary),
                ("Day Window", external_shaded_day_summary),
            ],
            "External Non-Shaded": [
                ("All Day", external_nonshaded_all_summary),
                ("Day Window", external_nonshaded_day_summary),
            ],
            "Internal Shaded": [
                ("All Day", internal_shaded_surface_all_summary),
                ("Day Window", internal_shaded_day_summary),
            ],
            "Internal Non-Shaded": [
                ("All Day", internal_nonshaded_surface_all_summary),
                ("Day Window", internal_nonshaded_day_summary),
            ],
            "Internal Floor": [
                ("All Day", internal_floor_all_summary),
                ("Day Window", internal_floor_day_summary),
            ],
        })

        st.download_button(
            "Download processed sheds surface Excel report",
            data=excel_data5,
            file_name=f"processed_sheds_stats_{start_date}_to_{end_date}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    if internal_allday_map is not None and sheds_allday_map is not None:
        st.subheader("Results")
        st.info("ΔT1 = TII - C, ΔT2 = TII - R, ΔT3 = C - R")

        delta_options = st.multiselect(
            "Select delta(s) to calculate",
            ["ΔT1", "ΔT2", "ΔT3"]
        )

        if delta_options:
            # Internal temperature deltas
            internal_delta_all_df = calculate_delta_summary(
                temp_all_summary,
                "Temperature Mean",
                delta_options
            )

            internal_delta_day_df = calculate_delta_summary(
                temp_day_summary,
                "Temperature Mean",
                delta_options
            )

            st.write("Internal Temperature Deltas")
            c1, c2 = st.columns(2)
            with c1:
                st.write("All day")
                st.dataframe(internal_delta_all_df, use_container_width=True)
            with c2:
                st.write("Day window")
                st.dataframe(internal_delta_day_df, use_container_width=True)

            # External shaded deltas
            external_shaded_delta_all_df = calculate_delta_summary(
                external_shaded_all_summary,
                "Ext Surface T Mean",
                delta_options
            )

            external_shaded_delta_day_df = calculate_delta_summary(
                external_shaded_day_summary,
                "Ext Surface T Mean",
                delta_options
            )

            st.write("External Shaded Surface Deltas")
            c1, c2 = st.columns(2)
            with c1:
                st.write("All day")
                st.dataframe(external_shaded_delta_all_df, use_container_width=True)
            with c2:
                st.write("Day window")
                st.dataframe(external_shaded_delta_day_df, use_container_width=True)

            # External  non-shaded deltas
            external_nonshaded_delta_all_df = calculate_delta_summary(
                external_nonshaded_all_summary,
                "Ext Surface T Mean",
                delta_options
            )

            external_nonshaded_delta_day_df = calculate_delta_summary(
                external_nonshaded_day_summary,
                "Ext Surface T Mean",
                delta_options
            )

            st.write("External Non-Shaded Surface Deltas")
            c1, c2 = st.columns(2)
            with c1:
                st.write("All day")
                st.dataframe(external_nonshaded_delta_all_df, use_container_width=True)
            with c2:
                st.write("Day window")
                st.dataframe(external_nonshaded_delta_day_df, use_container_width=True)

            # Internal shaded deltas
            internal_shaded_delta_all_df = calculate_delta_summary(
                internal_shaded_surface_all_summary,
                "Int Surface T Mean",
                delta_options
            )

            internal_shaded_delta_day_df = calculate_delta_summary(
                internal_shaded_day_summary,
                "Int Surface T Mean",
                delta_options
            )

            st.write("Internal Shaded Surface Deltas")
            c1, c2 = st.columns(2)
            with c1:
                st.write("All day")
                st.dataframe(internal_shaded_delta_all_df, use_container_width=True)
            with c2:
                st.write("Day window")
                st.dataframe(internal_shaded_delta_day_df, use_container_width=True)

            # Internal non-shaded deltas
            internal_nonshaded_delta_all_df = calculate_delta_summary(
                internal_nonshaded_surface_all_summary,
                "Int Surface T Mean",
                delta_options
            )

            internal_nonshaded_delta_day_df = calculate_delta_summary(
                internal_nonshaded_day_summary,
                "Int Surface T Mean",
                delta_options
            )

            st.write("Internal Non-Shaded Surface Deltas")
            c1, c2 = st.columns(2)
            with c1:
                st.write("All day")
                st.dataframe(internal_nonshaded_delta_all_df, use_container_width=True)
            with c2:
                st.write("Day window")
                st.dataframe(internal_nonshaded_delta_day_df, use_container_width=True)

            # Internal floor deltas
            internal_floor_delta_all_df = calculate_delta_summary(
                internal_floor_all_summary,
                "Int Floor T Mean",
                delta_options
            )     
            internal_floor_delta_day_df = calculate_delta_summary(
                internal_floor_day_summary,
                "Int Floor T Mean",
                delta_options
            )     

            st.write("Internal Floor Deltas")
            c1, c2 = st.columns(2)
            with c1:
                st.write("All day")
                st.dataframe(internal_floor_delta_all_df, use_container_width=True)
            with c2:
                st.write("Day window")
                st.dataframe(internal_floor_delta_day_df, use_container_width=True)

            # download deltas as Excel
            excel_data6 = export_excel_grouped({
                "Internal Deltas": [
                    ("All Day", internal_delta_all_df),
                    ("Day Window", internal_delta_day_df),
                ],

                "External Shaded": [
                    ("All Day", external_shaded_delta_all_df),
                    ("Day Window", external_shaded_delta_day_df),
                ],

                "External Non-Shaded": [
                    ("All Day", external_nonshaded_delta_all_df),
                    ("Day Window", external_nonshaded_delta_day_df),
                ],

                "Internal Shaded": [
                    ("All Day", internal_shaded_delta_all_df),
                    ("Day Window", internal_shaded_delta_day_df),
                ],

                "Internal Non-Shaded": [
                    ("All Day", internal_nonshaded_delta_all_df),
                    ("Day Window", internal_nonshaded_delta_day_df),
                ],

                "Internal Floor": [
                    ("All Day", internal_floor_delta_all_df),
                    ("Day Window", internal_floor_delta_day_df),
                ],
            })

            st.download_button(
                "Download deltas Excel report",
                data=excel_data6,
                file_name=f"deltas_stats_{start_date}_to_{end_date}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
